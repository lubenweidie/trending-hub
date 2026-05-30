"""一键发布入口：采集 → 生成 → 多平台发布 → Word 存档"""
import config_loader  # noqa: E402,F401
import os
import sys
import shutil
import subprocess
from pathlib import Path
from datetime import datetime

from article_utils import parse_article

HERE = Path(__file__).parent
OUTPUT_DIR = HERE / "output" / "articles"
LOGS_DIR = HERE / "logs"
DEFAULT_PLATFORM = "baijiahao"


class TeeLogger:
    """同时输出到终端和日志文件"""
    def __init__(self, log_path: Path):
        LOGS_DIR.mkdir(exist_ok=True)
        self.terminal = sys.stdout
        self.log = open(log_path, "w", encoding="utf-8")

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)

    def flush(self):
        self.terminal.flush()
        self.log.flush()

    def close(self):
        self.log.close()


def _publish_mode() -> str:
    """从 apikeys.conf 读取模式：1=草稿 2=发布（G5 DRY — 单一来源）"""
    num = int(os.environ.get("PUBLISH_MODE", "1"))
    return "publish" if num == 2 else "draft"


def run_pipeline(article_limit: int = 0, target_platforms: str = "",
                 platform_account_counts: str = ""):
    """执行内容管线，实时输出到终端和日志"""
    cmd = [sys.executable, str(HERE / "pipeline.py"),
           "--article-limit", str(article_limit)]
    if target_platforms:
        cmd += ["--target-platforms", target_platforms]
    if platform_account_counts:
        cmd += ["--platform-account-counts", platform_account_counts]
    proc = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        bufsize=1,
    )
    for line in proc.stdout:
        print(line, end="")
    proc.wait()
    return proc.returncode == 0


def export_word(article_path: Path):
    """将信息源原文导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("[Word] python-docx 未安装，跳过 Word 导出")
        return

    source_file = article_path.with_suffix(".source.txt")
    if not source_file.exists():
        print("[Word] 找不到源内容文件，跳过")
        return

    source_text = source_file.read_text(encoding="utf-8")
    if not source_text.strip():
        print("[Word] 源内容为空，跳过")
        return

    # 从 .md 获取标题和来源信息
    title = "无标题"
    source_info = ""
    if article_path.exists():
        md_text = article_path.read_text(encoding="utf-8")
        for line in md_text.strip().split("\n"):
            if line.startswith("# "):
                title = line[2:].strip()
            elif line.startswith("> 来源"):
                source_info = line.replace("> ", "").strip()

    doc = Document()
    doc.add_heading(title, level=1)

    if source_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(source_info)
        run.font.size = Pt(9)

    doc.add_heading("原始信息源内容", level=2)
    doc.add_paragraph(source_text)

    word_path = article_path.with_suffix(".docx")
    doc.save(str(word_path))
    print(f"[Word] 已保存源内容: {word_path.name}")

    # 清理临时文件
    source_file.unlink(missing_ok=True)


def _archive_article(article_file: Path, platform: str):
    """归档单篇文章到 published/"""
    from article_utils import archive_article
    archive_article(article_file, platform, OUTPUT_DIR)


def _total_account_count(platforms: list[str]) -> int:
    """计算总任务数：账号数 × 每账号篇数"""
    from services import account_rotator
    per = int(os.environ.get("ARTICLES_PER_ACCOUNT", "1"))
    total = 0
    for p in platforms:
        total += account_rotator.get_account_count(p)
    return total * per


def _per_platform_counts(platforms: list[str]) -> dict:
    """返回 {平台: 需要生成的篇数}，已含 ARTICLES_PER_ACCOUNT"""
    from services import account_rotator
    per = int(os.environ.get("ARTICLES_PER_ACCOUNT", "1"))
    return {p: account_rotator.get_account_count(p) * per for p in platforms}


def _publish_one(pub, article_file, plat, acc, title, mode, max_retries,
                 retry_delay, tracker, print_lock):
    """单次发布任务（含重试），返回 (plat, account_name, title, ok, error)"""
    import time
    account_name = acc["name"]
    tag = f"[{plat}][{account_name}]"
    last_error = ""
    for attempt in range(max_retries + 1):
        if attempt > 0:
            with print_lock:
                print(f"{tag} 重试 {attempt}/{max_retries}，等待 {retry_delay}s...")
            time.sleep(retry_delay)
        try:
            ok = pub.publish(article_path=str(article_file), mode=mode)
        except Exception as e:
            last_error = str(e)[:200]
            with print_lock:
                print(f"{tag} 异常: {last_error}")
            ok = False
        if ok:
            _finalize_publish_success(pub, article_file, plat, title, tracker, print_lock)
            with print_lock:
                print(f"\n{tag} ✓ 成功")
            return plat, account_name, title, True, ""
        last_error = f"发布失败（已重试{attempt + 1}次）"
    with print_lock:
        print(f"\n{tag} ✗ 失败（已重试 {max_retries} 次）")
    return plat, account_name, title, False, last_error


def _finalize_publish_success(pub, article_file: Path, plat: str, title: str, tracker, print_lock):
    """发布成功后处理：截图 + Word 导出 + 归档 + 追踪记录"""
    try:
        pub.screenshot(article_file.parent)
    except Exception as e:
        with print_lock:
            print(f"  截图异常: {e}")
    export_word(article_file)
    _archive_article(article_file, plat)
    tracker.add_entry(title, plat, article_file.parent.name)


def publish_to_platforms(platforms: list[str], mode: str = "draft", article_count: int = None):
    """并发多账号发布：每账号取1篇不同文章，同平台多账号并行"""
    import concurrent.futures
    import threading
    from publishers import REGISTRY
    from publishers.base import BasePublisher
    from services import account_rotator, notifier, tracker

    result = notifier.get_result()
    print_lock = threading.Lock()

    # Step 1: 计算任务数 + 取文章（按 profile 分组，减少 profile 切换次数）
    per_account = int(os.environ.get("ARTICLES_PER_ACCOUNT", "1"))
    all_accounts = []  # [(platform, account_dict)]
    for plat in platforms:
        if plat not in REGISTRY:
            print(f"[Skip] 未知平台: {plat}")
            continue
        for acc in account_rotator.get_all_accounts(plat):
            for _ in range(per_account):
                all_accounts.append((plat, acc))

    # 校验：每个 profile 缺失的平台仅跳过该平台，不影响其他平台
    profile_platforms: dict[str, set] = {}  # profile → {已覆盖平台}
    for plat, acc in all_accounts:
        pid = acc.get("opencli_profile", "")
        if pid not in profile_platforms:
            profile_platforms[pid] = set()
        profile_platforms[pid].add(plat)
    target_plat_set = set(platforms)
    for pid, pts in profile_platforms.items():
        missing = target_plat_set - pts
        if missing:
            print(f"[Warn] profile '{pid or '默认'}' 未覆盖: {', '.join(missing)}，跳过对应平台")
    # 未覆盖的平台不参与（all_accounts 只包含已配置的，无需额外过滤）

    # 按 opencli_profile 排序：同一 profile 的多平台任务相邻，避免重复切换
    all_accounts.sort(key=lambda x: x[1].get("opencli_profile", ""))

    total_tasks = len(all_accounts)
    article_files = BasePublisher.find_recent_articles(str(OUTPUT_DIR), count=total_tasks)

    if not article_files:
        print("[Exit] 没有文章")
        notifier.finish_result()
        return False

    max_retries = int(os.environ.get("RETRY_COUNT", "0"))
    retry_delay = int(os.environ.get("RETRY_DELAY", "60"))
    max_workers = int(os.environ.get("MAX_CONCURRENCY", "3"))

    print(f"\n{'='*55}")
    print(f"  并发发布: {len(platforms)}平台 × {total_tasks}账号（并行上限 {max_workers}）")
    print(f"  文章: {len(article_files)}篇 | 模式: {'发布' if mode == 'publish' else '草稿'}")
    for plat in platforms:
        accs = account_rotator.get_all_accounts(plat)
        names = ", ".join(a["name"] for a in accs)
        print(f"  [{plat}] {len(accs)}个账号: {names}")
    print(f"{'='*55}")

    # 自动打开所有 Chrome Profile 窗口（确保 opencli 可连接）
    account_rotator.ensure_chrome_windows()

    # Step 2: 串行准备 — 切 profile + 打开编辑器（绑定 session 到对应 Chrome profile）
    tasks = []
    for idx, (plat, acc) in enumerate(all_accounts):
        if idx >= len(article_files):
            break

        account_rotator.switch_profile(acc)
        pub_cls = REGISTRY[plat]
        pub = pub_cls(account_suffix=acc["name"], profile_id=acc.get("opencli_profile", ""))
        pub._account_name = acc["name"]

        article_file = article_files[idx]
        article = parse_article(article_file)

        tag = f"[{plat}][{acc['name']}]"
        if not pub.prepare_editor(article["title"]):
            with print_lock:
                print(f"{tag} 编辑器准备失败，跳过")
            continue

        with print_lock:
            print(f"\n{tag} 准备就绪 → {article['title'][:40]}")
        tasks.append((pub, article_file, plat, acc, article["title"]))

    # Step 3: 并发执行
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(_publish_one, pub, af, plat, acc, atitle,
                                   mode, max_retries, retry_delay, tracker, print_lock)
                   for pub, af, plat, acc, atitle in tasks]
        for future in concurrent.futures.as_completed(futures):
            plat, account_name, title, ok, error = future.result()
            result.add_platform(f"{plat}:{account_name}", title, ok, error=error)

    # Step 4: 清理剩余文章
    used = len(tasks)
    leftovers = article_files[used:]
    if leftovers:
        leftover_dir = OUTPUT_DIR / "_剩余"
        leftover_dir.mkdir(exist_ok=True)
        for f in leftovers:
            src = f.parent
            dst = leftover_dir / src.name
            if dst.exists():
                shutil.rmtree(str(dst))
            shutil.move(str(src), str(dst))
            print(f"[剩余] {src.name}")

    result.set_articles(len(article_files), len(leftovers))
    notifier.finish_result()

    ok_count = sum(1 for p in result.platforms.values() if p["status"] == "ok")
    return ok_count > 0


def main():
    import argparse
    parser = argparse.ArgumentParser(description="一键发布到多平台")
    parser.add_argument("-p", "--platforms", default="",
                        help="目标平台（逗号分隔，如: baijiahao,toutiao）")
    parser.add_argument("--publish", action="store_true",
                        help="立即发布（默认存草稿）")
    args = parser.parse_args()

    # 日志：同时输出到终端和文件
    LOGS_DIR.mkdir(exist_ok=True)
    log_ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = LOGS_DIR / f"publish_{log_ts}.log"
    # 清理旧日志，保留最近 30 个
    old_logs = sorted(LOGS_DIR.glob("publish_*.log"), key=lambda p: p.stat().st_mtime, reverse=True)
    for f in old_logs[29:]:
        f.unlink()
    logger = TeeLogger(log_path)
    sys.stdout = logger
    sys.stderr = logger
    try:
        _main_impl(args)
    finally:
        logger.close()
        sys.stdout = logger.terminal
        sys.stderr = logger.terminal


def _main_impl(args):
    # 发布模式：--publish 标志 > 环境变量 PUBLISH_MODE=2 > apikeys.conf
    if args.publish:
        mode = "publish"
    else:
        mode = _publish_mode()

    # 平台列表
    platforms_env = os.environ.get("PUBLISH_PLATFORMS", "")
    plat_str = args.platforms or platforms_env or DEFAULT_PLATFORM
    platforms = [p.strip() for p in plat_str.split(",") if p.strip()]

    # 文章数量 = 总账号数 × ARTICLES_PER_ACCOUNT
    per_account = int(os.environ.get("ARTICLES_PER_ACCOUNT", "1"))
    article_count = _total_account_count(platforms)
    counts = _per_platform_counts(platforms)
    counts_str = ",".join(f"{k}:{v}" for k, v in counts.items())

    print("=" * 55)
    print(f"  一键发布 → {', '.join(platforms)}")
    print(f"  模式: {'立即发布' if mode == 'publish' else '存草稿'}")
    count_detail = " + ".join(f"{p}×{counts.get(p, 1)}" for p in platforms)
    extra = f"（每账号 {per_account} 篇）" if per_account > 1 else ""
    print(f"  账号: {count_detail}，生成 {article_count} 篇{extra}")
    print("=" * 55)

    # Step 1: 采集 + 生成
    print(f"\n[1/2] 采集热点 + AI摘要 + 生成文章（{article_count}篇）...")
    if not run_pipeline(article_limit=article_count, target_platforms=plat_str,
                        platform_account_counts=counts_str):
        print("[FAIL] 管线执行失败")
        sys.exit(1)

    # Step 2: 多平台发布
    print(f"\n[2/2] 发布到 {len(platforms)} 个平台...")
    publish_to_platforms(platforms, mode, article_count=article_count)

    from publishers.base import BasePublisher
    BasePublisher.stop_daemon()

    print("\n[OK] 全部完成")


if __name__ == "__main__":
    main()
