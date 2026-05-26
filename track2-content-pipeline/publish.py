"""一键发布入口：采集 → 生成 → 多平台发布 → Word 存档"""
import config_loader  # noqa: E402,F401
import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime

HERE = Path(__file__).parent
OUTPUT_DIR = HERE / "output" / "articles"
DEFAULT_PLATFORM = "baijiahao"


def _publish_mode() -> str:
    """从 apikeys.conf 读取模式：1=草稿 2=发布（G5 DRY — 单一来源）"""
    num = int(os.environ.get("PUBLISH_MODE", "1"))
    return "publish" if num == 2 else "draft"


def run_pipeline():
    """执行内容管线"""
    result = subprocess.run(
        [sys.executable, str(HERE / "pipeline.py")],
        capture_output=False
    )
    return result.returncode == 0


def export_word(article_path: Path):
    """将信息源原文导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("[Word] python-docx 未安装，跳过 Word 导出")
        return

    source_file = article_path.with_suffix(".source.txt")
    if not source_file.exists():
        print("[Word] 找不到源内容文件，跳过")
        return

    source_text = source_file.read_text(encoding="utf-8")
    if not source_text.strip():
        print("[Word] 源内容为空，跳过")
        return

    # 从 .md 获取标题和来源信息
    title = "无标题"
    source_info = ""
    if article_path.exists():
        md_text = article_path.read_text(encoding="utf-8")
        for line in md_text.strip().split("\n"):
            if line.startswith("# "):
                title = line[2:].strip()
            elif line.startswith("> 来源"):
                source_info = line.replace("> ", "").strip()

    doc = Document()
    doc.add_heading(title, level=1)

    if source_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(source_info)
        run.font.size = Pt(9)

    doc.add_heading("原始信息源内容", level=2)
    doc.add_paragraph(source_text)

    word_dir = OUTPUT_DIR / "published"
    word_dir.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    word_path = word_dir / f"{ts}_{Path(article_path).stem}.docx"
    doc.save(str(word_path))
    print(f"[Word] 已保存源内容: {word_path.name}")

    # 清理临时文件
    source_file.unlink(missing_ok=True)


def _archive_article(article_file: Path, platform: str):
    """归档单篇文章到 published/，追加平台后缀"""
    import shutil
    article_dir = article_file.parent
    new_name = f"{article_dir.name}-{platform}"
    new_path = article_dir.parent / new_name
    article_dir.rename(new_path)
    article_dir = new_path
    published_dir = OUTPUT_DIR / "published"
    published_dir.mkdir(exist_ok=True)
    archived = published_dir / article_dir.name
    if archived.exists():
        shutil.rmtree(str(archived))
    shutil.move(str(article_dir), str(archived))
    print(f"[归档] {archived.name}")


def publish_to_platforms(platforms: list[str], mode: str = "draft", article_count: int = None):
    """发布到指定平台：每平台取不同文章"""
    from publishers import REGISTRY
    from publishers.base import BasePublisher

    count = article_count or len(platforms)
    article_files = BasePublisher.find_recent_articles(str(OUTPUT_DIR), count=count)

    if not article_files:
        print("[Exit] 没有文章")
        return False

    print(f"\n找到 {len(article_files)} 篇文章，目标 {len(platforms)} 个平台")

    success = True
    for i, name in enumerate(platforms):
        if i >= len(article_files):
            print(f"[Skip] 文章不够，{name} 跳过")
            continue
        if name not in REGISTRY:
            print(f"[Skip] 未知平台: {name}")
            continue

        article_file = article_files[i]
        article = BasePublisher.parse_article(article_file)
        print(f"\n[{name}] → {article['title'][:40]}")

        pub_cls = REGISTRY[name]
        pub = pub_cls()
        if pub.publish(article_path=str(article_file), mode=mode):
            export_word(article_file)
            _archive_article(article_file, name)
        else:
            success = False

    return success


def main():
    import argparse
    parser = argparse.ArgumentParser(description="一键发布到多平台")
    parser.add_argument("-p", "--platforms", default="",
                        help="目标平台（逗号分隔，如: baijiahao,toutiao）")
    parser.add_argument("--publish", action="store_true",
                        help="立即发布（默认存草稿）")
    args = parser.parse_args()

    # 发布模式：--publish 标志 > 环境变量 PUBLISH_MODE=2 > apikeys.conf
    if args.publish:
        mode = "publish"
    else:
        mode = _publish_mode()

    # 平台列表
    platforms_env = os.environ.get("PUBLISH_PLATFORMS", "")
    plat_str = args.platforms or platforms_env or DEFAULT_PLATFORM
    platforms = [p.strip() for p in plat_str.split(",") if p.strip()]

    # 文章数量：从 apikeys.conf 的 ARTICLE_LIMIT 读取
    article_count = int(os.environ.get("ARTICLE_LIMIT", "1"))

    print("=" * 55)
    print(f"  一键发布 → {', '.join(platforms)}")
    print(f"  模式: {'立即发布' if mode == 'publish' else '存草稿'}")
    print(f"  生成: {article_count} 篇")
    print("=" * 55)

    # Step 1: 采集 + 生成
    print(f"\n[1/2] 采集热点 + AI摘要 + 生成文章（{article_count}篇）...")
    if not run_pipeline():
        print("[FAIL] 管线执行失败")
        sys.exit(1)

    # Step 2: 多平台发布
    print(f"\n[2/2] 发布到 {len(platforms)} 个平台...")
    publish_to_platforms(platforms, mode, article_count=article_count)

    print("\n[OK] 全部完成")


if __name__ == "__main__":
    main()
